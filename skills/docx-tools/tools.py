"""
DOCX Tools Skill

Word document toolkit using python-docx for reading, creating, and manipulating Word documents.
Supports professional formatting with tables, colors, and beautiful styling.
"""
import os
import logging
import re
from typing import Dict, Any, List, Optional, Tuple

from Jotty.core.utils.skill_status import SkillStatus
from Jotty.core.utils.tool_helpers import tool_response, tool_error, tool_wrapper

logger = logging.getLogger(__name__)

# Professional color scheme (RGB values)

# Status emitter for progress updates
status = SkillStatus("docx-tools")

COLORS = {
    'primary': (41, 65, 114),      # Deep blue - headers
    'secondary': (89, 89, 89),     # Dark gray - text
    'accent': (0, 112, 192),       # Bright blue - links/highlights
    'success': (84, 130, 53),      # Green - checked items
    'light_bg': (242, 242, 242),   # Light gray - table header bg
    'border': (191, 191, 191),     # Border gray
    'white': (255, 255, 255),
}


@tool_wrapper()
def read_docx_tool(params: Dict[str, Any]) -> Dict[str, Any]:
    """
    Read text content from a Word document.

    Args:
        params: Dictionary containing:
            - file_path (str, required): Path to the .docx file
            - include_tables (bool, optional): Include table content (default: True)

    Returns:
        Dictionary with:
            - success (bool): Whether reading succeeded
            - text (str): Extracted text content
            - paragraphs (list): List of paragraph texts
            - tables (list): List of tables (if include_tables is True)
            - error (str, optional): Error message if failed
    """
    status.set_callback(params.pop('_status_callback', None))

    try:
        from docx import Document
    except ImportError:
        return {
            'success': False,
            'error': 'python-docx not installed. Install with: pip install python-docx'
        }

    file_path = params.get('file_path')
    include_tables = params.get('include_tables', True)

    if not file_path:
        return {'success': False, 'error': 'file_path parameter is required'}

    if not os.path.exists(file_path):
        return {'success': False, 'error': f'File not found: {file_path}'}

    try:
        doc = Document(file_path)

        paragraphs = [para.text for para in doc.paragraphs]
        full_text = '\n'.join(paragraphs)

        tables_data = []
        if include_tables:
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_rows.append(row_data)
                tables_data.append(table_rows)

        logger.info(f"Document read: {file_path}")

        return {
            'success': True,
            'text': full_text,
            'paragraphs': paragraphs,
            'tables': tables_data,
            'paragraph_count': len(paragraphs),
            'table_count': len(tables_data)
        }

    except Exception as e:
        logger.error(f"Failed to read document: {e}", exc_info=True)
        return {'success': False, 'error': f'Failed to read document: {str(e)}'}


@tool_wrapper()
def create_docx_tool(params: Dict[str, Any]) -> Dict[str, Any]:
    """
    Create a new Word document with proper markdown formatting.

    Supports markdown conversion:
    - # Header → Heading 1
    - ## Header → Heading 2
    - ### Header → Heading 3
    - - [ ] item → Checkbox (unchecked)
    - - [x] item → Checkbox (checked)
    - - item or * item → Bullet point
    - 1. item → Numbered list
    - **bold** → Bold text
    - *italic* → Italic text

    Args:
        params: Dictionary containing:
            - content (str or list, required): Markdown content or list of paragraphs
            - output_path (str, required): Output file path
            - title (str, optional): Document title (added as Heading1)

    Returns:
        Dictionary with:
            - success (bool): Whether creation succeeded
            - file_path (str): Path to created document
            - error (str, optional): Error message if failed
    """
    status.set_callback(params.pop('_status_callback', None))

    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {
            'success': False,
            'error': 'python-docx not installed. Install with: pip install python-docx'
        }

    import re

    content = params.get('content')
    output_path = params.get('output_path')
    title = params.get('title')

    if not content:
        return {'success': False, 'error': 'content parameter is required'}

    if not output_path:
        return {'success': False, 'error': 'output_path parameter is required'}

    try:
        doc = Document()

        if title:
            doc.add_heading(title, level=0)

        if isinstance(content, str):
            lines = content.split('\n')
        elif isinstance(content, list):
            lines = content
        else:
            lines = [str(content)]

        for line in lines:
            line_stripped = line.strip()

            if not line_stripped:
                # Empty line - add spacing
                doc.add_paragraph()
                continue

            # Heading 1: # Header
            if line_stripped.startswith('# '):
                text = line_stripped[2:].strip()
                doc.add_heading(text, level=1)

            # Heading 2: ## Header
            elif line_stripped.startswith('## '):
                text = line_stripped[3:].strip()
                doc.add_heading(text, level=2)

            # Heading 3: ### Header
            elif line_stripped.startswith('### '):
                text = line_stripped[4:].strip()
                doc.add_heading(text, level=3)

            # Heading 4: #### Header
            elif line_stripped.startswith('#### '):
                text = line_stripped[5:].strip()
                doc.add_heading(text, level=4)

            # Checkbox unchecked: - [ ] item
            elif line_stripped.startswith('- [ ] ') or line_stripped.startswith('* [ ] '):
                text = line_stripped[6:].strip()
                para = doc.add_paragraph(style='List Bullet')
                para.add_run('☐ ' + text)

            # Checkbox checked: - [x] item
            elif line_stripped.startswith('- [x] ') or line_stripped.startswith('- [X] ') or \
                 line_stripped.startswith('* [x] ') or line_stripped.startswith('* [X] '):
                text = line_stripped[6:].strip()
                para = doc.add_paragraph(style='List Bullet')
                para.add_run('☑ ' + text)

            # Bullet point: - item or * item
            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                text = line_stripped[2:].strip()
                para = doc.add_paragraph(style='List Bullet')
                _add_formatted_text(para, text)

            # Numbered list: 1. item
            elif re.match(r'^\d+\.\s', line_stripped):
                text = re.sub(r'^\d+\.\s*', '', line_stripped)
                para = doc.add_paragraph(style='List Number')
                _add_formatted_text(para, text)

            # Horizontal rule: --- or ***
            elif line_stripped in ['---', '***', '___']:
                # Add a subtle separator
                para = doc.add_paragraph('─' * 50)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Regular paragraph with inline formatting
            else:
                para = doc.add_paragraph()
                _add_formatted_text(para, line_stripped)

        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        doc.save(output_path)
        logger.info(f"Document created: {output_path}")

        return {
            'success': True,
            'file_path': output_path,
            'line_count': len(lines)
        }

    except Exception as e:
        logger.error(f"Failed to create document: {e}", exc_info=True)
        return {'success': False, 'error': f'Failed to create document: {str(e)}'}


def _add_formatted_text(paragraph, text: str):
    """
    Add text to paragraph with inline markdown formatting.

    Handles:
    - **bold** → Bold
    - *italic* → Italic
    - `code` → Monospace (simulated)
    """
    import re

    # Pattern for bold, italic, code
    # Process in order: bold first (** **), then italic (* *)
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before this match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        full_match = match.group(0)

        if full_match.startswith('**') and full_match.endswith('**'):
            # Bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif full_match.startswith('`') and full_match.endswith('`'):
            # Code - use different font
            run = paragraph.add_run(match.group(4))
            run.font.name = 'Courier New'
        elif full_match.startswith('*') and full_match.endswith('*'):
            # Italic
            run = paragraph.add_run(match.group(3))
            run.italic = True

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


@tool_wrapper()
def create_professional_checklist_tool(params: Dict[str, Any]) -> Dict[str, Any]:
    """
    Create a beautifully formatted professional checklist document.

    Features:
    - Professional color scheme (deep blue headers, clean styling)
    - Table-based checklist with columns: Item, Reference, Status
    - Form fields at top (Name, Date, Reviewer, Period)
    - Section hierarchy (Parts, numbered sections)
    - Checkboxes with proper formatting

    Args:
        params: Dictionary containing:
            - content (str, required): Markdown content with checklist items
            - output_path (str, required): Output file path
            - title (str, optional): Main title
            - subtitle (str, optional): Subtitle
            - organization (str, optional): Organization name for header
            - include_form_fields (bool, optional): Add form fields at top (default: True)

    Returns:
        Dictionary with success status and file path
    """
    status.set_callback(params.pop('_status_callback', None))

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return {
            'success': False,
            'error': 'python-docx not installed. Install with: pip install python-docx'
        }

    content = params.get('content', '')
    output_path = params.get('output_path')
    title = params.get('title', 'Checklist')
    subtitle = params.get('subtitle', '')
    organization = params.get('organization', '')
    include_form_fields = params.get('include_form_fields', True)

    if not output_path:
        return {'success': False, 'error': 'output_path parameter is required'}

    try:
        doc = Document()

        # Set document margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # === TITLE SECTION ===
        _add_title_section(doc, title, subtitle, organization)

        # === FORM FIELDS ===
        if include_form_fields:
            _add_form_fields(doc)

        # === PARSE AND ADD CONTENT ===
        _parse_and_add_content(doc, content)

        # Save document
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        doc.save(output_path)
        logger.info(f"Professional checklist created: {output_path}")

        return {
            'success': True,
            'file_path': output_path
        }

    except Exception as e:
        logger.error(f"Failed to create professional checklist: {e}", exc_info=True)
        return {'success': False, 'error': str(e)}


def _add_title_section(doc, title: str, subtitle: str, organization: str):
    """Add professional title section with colors."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Main title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title.upper())
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(*COLORS['primary'])
    title_run.font.name = 'Calibri'

    # Subtitle
    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(subtitle)
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = RGBColor(*COLORS['secondary'])
        sub_run.font.name = 'Calibri'

    # Organization
    if organization:
        org_para = doc.add_paragraph()
        org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        org_run = org_para.add_run(organization)
        org_run.font.size = Pt(11)
        org_run.font.color.rgb = RGBColor(*COLORS['accent'])
        org_run.font.italic = True

    # Divider line
    div_para = doc.add_paragraph()
    div_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_run = div_para.add_run('━' * 60)
    div_run.font.color.rgb = RGBColor(*COLORS['primary'])

    doc.add_paragraph()  # Spacing


def _add_form_fields(doc):
    """Add form fields section at top of document."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT

    # Create a 2x4 table for form fields
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    fields = [
        ('Entity/Fund Name:', ''),
        ('Review Period:', ''),
        ('Reviewer:', ''),
        ('Date:', ''),
    ]

    for i, (label, _) in enumerate(fields):
        row = i // 2
        col = (i % 2) * 2

        # Label cell
        label_cell = table.rows[row].cells[col]
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(10)
        label_run.font.color.rgb = RGBColor(*COLORS['secondary'])

        # Value cell (underline for filling)
        value_cell = table.rows[row].cells[col + 1]
        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run('_' * 25)
        value_run.font.size = Pt(10)
        value_run.font.color.rgb = RGBColor(*COLORS['border'])

    doc.add_paragraph()  # Spacing


def _parse_and_add_content(doc, content: str):
    """Parse markdown content and add formatted sections."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    lines = content.split('\n')
    current_section = None
    current_items = []

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Part header (e.g., "PART A:" or "## PART A")
        if line.upper().startswith('PART ') or (line.startswith('## ') and 'PART' in line.upper()):
            # Flush previous section
            if current_items:
                _add_checklist_table(doc, current_section, current_items)
                current_items = []

            # Add part header
            clean_title = re.sub(r'^##?\s*', '', line)
            _add_part_header(doc, clean_title)
            current_section = clean_title
            i += 1
            continue

        # Section header (e.g., "# Title" or "## 1. Section")
        if line.startswith('# ') or line.startswith('## '):
            # Flush previous items
            if current_items:
                _add_checklist_table(doc, current_section, current_items)
                current_items = []

            clean_title = re.sub(r'^##?\s*', '', line)

            if line.startswith('# '):
                _add_section_header(doc, clean_title, level=1)
            else:
                _add_section_header(doc, clean_title, level=2)

            current_section = clean_title
            i += 1
            continue

        # Subsection header (### )
        if line.startswith('### '):
            if current_items:
                _add_checklist_table(doc, current_section, current_items)
                current_items = []

            clean_title = line[4:].strip()
            _add_section_header(doc, clean_title, level=3)
            current_section = clean_title
            i += 1
            continue

        # Checklist item with checkbox
        if line.startswith('- [ ]') or line.startswith('- [x]') or line.startswith('- [X]'):
            checked = '[x]' in line.lower()
            item_text = re.sub(r'^-\s*\[[xX ]?\]\s*', '', line)

            # Check if there's a reference (text after | or in parentheses)
            reference = ''
            if '|' in item_text:
                parts = item_text.split('|', 1)
                item_text = parts[0].strip()
                reference = parts[1].strip()
            elif re.search(r'\(([^)]+)\)\s*$', item_text):
                match = re.search(r'\(([^)]+)\)\s*$', item_text)
                reference = match.group(1)
                item_text = item_text[:match.start()].strip()

            current_items.append({
                'text': item_text,
                'checked': checked,
                'reference': reference
            })
            i += 1
            continue

        # Regular bullet point (treat as checklist item)
        if line.startswith('- ') or line.startswith('* '):
            item_text = line[2:].strip()
            reference = ''
            if '|' in item_text:
                parts = item_text.split('|', 1)
                item_text = parts[0].strip()
                reference = parts[1].strip()

            current_items.append({
                'text': item_text,
                'checked': False,
                'reference': reference
            })
            i += 1
            continue

        # Regular paragraph
        if current_items:
            _add_checklist_table(doc, current_section, current_items)
            current_items = []

        para = doc.add_paragraph()
        _add_formatted_text(para, line)
        i += 1

    # Flush remaining items
    if current_items:
        _add_checklist_table(doc, current_section, current_items)


def _add_part_header(doc, title: str):
    """Add a PART header with professional styling."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_paragraph()  # Spacing

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = para.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(*COLORS['primary'])
    run.font.name = 'Calibri'

    # Add colored underline effect
    underline_para = doc.add_paragraph()
    underline_run = underline_para.add_run('━' * 40)
    underline_run.font.color.rgb = RGBColor(*COLORS['accent'])
    underline_run.font.size = Pt(8)


def _add_section_header(doc, title: str, level: int = 2):
    """Add section header with appropriate styling."""
    from docx.shared import Pt, RGBColor

    sizes = {1: 16, 2: 13, 3: 11}
    colors = {1: COLORS['primary'], 2: COLORS['secondary'], 3: COLORS['accent']}

    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = RGBColor(*colors.get(level, COLORS['secondary']))
    run.font.name = 'Calibri'


def _add_checklist_table(doc, section_title: str, items: List[Dict]):
    """Add a professional checklist table."""
    from docx.shared import Pt, RGBColor, Inches, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not items:
        return

    # Check if any items have references
    has_references = any(item.get('reference') for item in items)

    # Create table
    if has_references:
        table = doc.add_table(rows=len(items) + 1, cols=4)
        col_widths = [Inches(0.4), Inches(4.0), Inches(1.5), Inches(0.8)]
        headers = ['', 'Checklist Item', 'Reference', 'Status']
    else:
        table = doc.add_table(rows=len(items) + 1, cols=3)
        col_widths = [Inches(0.4), Inches(5.0), Inches(1.0)]
        headers = ['', 'Checklist Item', 'Status']

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style table
    table.style = 'Table Grid'

    # Set column widths
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = width

    # Header row
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*COLORS['white'])
        run.font.name = 'Calibri'

        # Set cell background color
        _set_cell_bg_color(cell, COLORS['primary'])

    # Data rows
    for row_idx, item in enumerate(items):
        row = table.rows[row_idx + 1]

        # Checkbox column
        checkbox_cell = row.cells[0]
        checkbox_para = checkbox_cell.paragraphs[0]
        checkbox_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if item.get('checked'):
            checkbox_run = checkbox_para.add_run('☑')
            checkbox_run.font.color.rgb = RGBColor(*COLORS['success'])
        else:
            checkbox_run = checkbox_para.add_run('☐')
            checkbox_run.font.color.rgb = RGBColor(*COLORS['secondary'])
        checkbox_run.font.size = Pt(12)

        # Item text column
        text_cell = row.cells[1]
        text_para = text_cell.paragraphs[0]
        text_run = text_para.add_run(item.get('text', ''))
        text_run.font.size = Pt(10)
        text_run.font.color.rgb = RGBColor(*COLORS['secondary'])
        text_run.font.name = 'Calibri'

        # Reference column (if applicable)
        if has_references:
            ref_cell = row.cells[2]
            ref_para = ref_cell.paragraphs[0]
            ref_run = ref_para.add_run(item.get('reference', ''))
            ref_run.font.size = Pt(9)
            ref_run.font.color.rgb = RGBColor(*COLORS['accent'])
            ref_run.font.italic = True

            status_cell = row.cells[3]
        else:
            status_cell = row.cells[2]

        # Status column (empty for user to fill)
        status_para = status_cell.paragraphs[0]
        status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Alternate row colors
        if row_idx % 2 == 1:
            for cell in row.cells:
                _set_cell_bg_color(cell, COLORS['light_bg'])

    doc.add_paragraph()  # Spacing after table


def _set_cell_bg_color(cell, color: Tuple[int, int, int]):
    """Set background color of a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(*color))
    cell._tc.get_or_add_tcPr().append(shading)


@tool_wrapper()
def add_paragraph_tool(params: Dict[str, Any]) -> Dict[str, Any]:
    """
    Add a paragraph to an existing Word document.

    Args:
        params: Dictionary containing:
            - file_path (str, required): Path to the .docx file
            - text (str, required): Paragraph text to add
            - style (str, optional): Paragraph style (Normal, Heading1, Heading2, etc.)

    Returns:
        Dictionary with:
            - success (bool): Whether addition succeeded
            - file_path (str): Path to modified document
            - error (str, optional): Error message if failed
    """
    status.set_callback(params.pop('_status_callback', None))

    try:
        from docx import Document
    except ImportError:
        return {
            'success': False,
            'error': 'python-docx not installed. Install with: pip install python-docx'
        }

    file_path = params.get('file_path')
    text = params.get('text')
    style = params.get('style', 'Normal')

    if not file_path:
        return {'success': False, 'error': 'file_path parameter is required'}

    if not text:
        return {'success': False, 'error': 'text parameter is required'}

    if not os.path.exists(file_path):
        return {'success': False, 'error': f'File not found: {file_path}'}

    try:
        doc = Document(file_path)

        if style.startswith('Heading'):
            level = int(style[-1]) if style[-1].isdigit() else 1
            doc.add_heading(text, level=level)
        else:
            para = doc.add_paragraph(text)
            try:
                para.style = style
            except KeyError:
                para.style = 'Normal'
                logger.warning(f"Style '{style}' not found, using 'Normal'")

        doc.save(file_path)
        logger.info(f"Paragraph added to: {file_path}")

        return {
            'success': True,
            'file_path': file_path,
            'style_applied': style
        }

    except Exception as e:
        logger.error(f"Failed to add paragraph: {e}", exc_info=True)
        return {'success': False, 'error': f'Failed to add paragraph: {str(e)}'}


@tool_wrapper()
def add_table_tool(params: Dict[str, Any]) -> Dict[str, Any]:
    """
    Add a table to an existing Word document.

    Args:
        params: Dictionary containing:
            - file_path (str, required): Path to the .docx file
            - data (list, required): List of lists representing table rows
            - headers (list, optional): List of header strings (if provided, adds header row)

    Returns:
        Dictionary with:
            - success (bool): Whether addition succeeded
            - file_path (str): Path to modified document
            - rows (int): Number of rows added
            - cols (int): Number of columns
            - error (str, optional): Error message if failed
    """
    status.set_callback(params.pop('_status_callback', None))

    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        return {
            'success': False,
            'error': 'python-docx not installed. Install with: pip install python-docx'
        }

    file_path = params.get('file_path')
    data = params.get('data')
    headers = params.get('headers')

    if not file_path:
        return {'success': False, 'error': 'file_path parameter is required'}

    if not data:
        return {'success': False, 'error': 'data parameter is required'}

    if not isinstance(data, list) or not all(isinstance(row, list) for row in data):
        return {'success': False, 'error': 'data must be a list of lists'}

    if not os.path.exists(file_path):
        return {'success': False, 'error': f'File not found: {file_path}'}

    try:
        doc = Document(file_path)

        all_rows = []
        if headers:
            all_rows.append(headers)
        all_rows.extend(data)

        if not all_rows:
            return {'success': False, 'error': 'No data to add'}

        num_cols = max(len(row) for row in all_rows)
        num_rows = len(all_rows)

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        for i, row_data in enumerate(all_rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    row.cells[j].text = str(cell_text)

        if headers:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        doc.save(file_path)
        logger.info(f"Table added to: {file_path}")

        return {
            'success': True,
            'file_path': file_path,
            'rows': num_rows,
            'cols': num_cols
        }

    except Exception as e:
        logger.error(f"Failed to add table: {e}", exc_info=True)
        return {'success': False, 'error': f'Failed to add table: {str(e)}'}


@tool_wrapper()
def add_image_tool(params: Dict[str, Any]) -> Dict[str, Any]:
    """
    Add an image to an existing Word document.

    Args:
        params: Dictionary containing:
            - file_path (str, required): Path to the .docx file
            - image_path (str, required): Path to the image file
            - width (float, optional): Image width in inches (default: 6.0)

    Returns:
        Dictionary with:
            - success (bool): Whether addition succeeded
            - file_path (str): Path to modified document
            - image_path (str): Path to the image that was added
            - error (str, optional): Error message if failed
    """
    status.set_callback(params.pop('_status_callback', None))

    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        return {
            'success': False,
            'error': 'python-docx not installed. Install with: pip install python-docx'
        }

    file_path = params.get('file_path')
    image_path = params.get('image_path')
    width = params.get('width', 6.0)

    if not file_path:
        return {'success': False, 'error': 'file_path parameter is required'}

    if not image_path:
        return {'success': False, 'error': 'image_path parameter is required'}

    if not os.path.exists(file_path):
        return {'success': False, 'error': f'Document not found: {file_path}'}

    if not os.path.exists(image_path):
        return {'success': False, 'error': f'Image not found: {image_path}'}

    try:
        doc = Document(file_path)

        doc.add_picture(image_path, width=Inches(width))

        doc.save(file_path)
        logger.info(f"Image added to: {file_path}")

        return {
            'success': True,
            'file_path': file_path,
            'image_path': image_path,
            'width_inches': width
        }

    except Exception as e:
        logger.error(f"Failed to add image: {e}", exc_info=True)
        return {'success': False, 'error': f'Failed to add image: {str(e)}'}


@tool_wrapper()
def replace_text_tool(params: Dict[str, Any]) -> Dict[str, Any]:
    """
    Find and replace text in a Word document.

    Args:
        params: Dictionary containing:
            - file_path (str, required): Path to the .docx file
            - find (str, required): Text to find
            - replace (str, required): Text to replace with

    Returns:
        Dictionary with:
            - success (bool): Whether replacement succeeded
            - file_path (str): Path to modified document
            - replacements (int): Number of replacements made
            - error (str, optional): Error message if failed
    """
    status.set_callback(params.pop('_status_callback', None))

    try:
        from docx import Document
    except ImportError:
        return {
            'success': False,
            'error': 'python-docx not installed. Install with: pip install python-docx'
        }

    file_path = params.get('file_path')
    find_text = params.get('find')
    replace_text = params.get('replace')

    if not file_path:
        return {'success': False, 'error': 'file_path parameter is required'}

    if not find_text:
        return {'success': False, 'error': 'find parameter is required'}

    if replace_text is None:
        return {'success': False, 'error': 'replace parameter is required'}

    if not os.path.exists(file_path):
        return {'success': False, 'error': f'File not found: {file_path}'}

    try:
        doc = Document(file_path)
        replacements = 0

        for para in doc.paragraphs:
            if find_text in para.text:
                for run in para.runs:
                    if find_text in run.text:
                        count = run.text.count(find_text)
                        run.text = run.text.replace(find_text, replace_text)
                        replacements += count

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if find_text in para.text:
                            for run in para.runs:
                                if find_text in run.text:
                                    count = run.text.count(find_text)
                                    run.text = run.text.replace(find_text, replace_text)
                                    replacements += count

        doc.save(file_path)
        logger.info(f"Text replaced in: {file_path}, {replacements} replacements")

        return {
            'success': True,
            'file_path': file_path,
            'replacements': replacements,
            'find': find_text,
            'replace': replace_text
        }

    except Exception as e:
        logger.error(f"Failed to replace text: {e}", exc_info=True)
        return {'success': False, 'error': f'Failed to replace text: {str(e)}'}


@tool_wrapper()
def get_styles_tool(params: Dict[str, Any]) -> Dict[str, Any]:
    """
    List available styles in a Word document.

    Args:
        params: Dictionary containing:
            - file_path (str, required): Path to the .docx file

    Returns:
        Dictionary with:
            - success (bool): Whether listing succeeded
            - styles (list): List of style names with their types
            - paragraph_styles (list): List of paragraph style names
            - character_styles (list): List of character style names
            - error (str, optional): Error message if failed
    """
    status.set_callback(params.pop('_status_callback', None))

    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        return {
            'success': False,
            'error': 'python-docx not installed. Install with: pip install python-docx'
        }

    file_path = params.get('file_path')

    if not file_path:
        return {'success': False, 'error': 'file_path parameter is required'}

    if not os.path.exists(file_path):
        return {'success': False, 'error': f'File not found: {file_path}'}

    try:
        doc = Document(file_path)

        all_styles = []
        paragraph_styles = []
        character_styles = []
        table_styles = []

        for style in doc.styles:
            style_info = {
                'name': style.name,
                'type': str(style.type).split('.')[-1].replace('>', '')
            }
            all_styles.append(style_info)

            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                paragraph_styles.append(style.name)
            elif style.type == WD_STYLE_TYPE.CHARACTER:
                character_styles.append(style.name)
            elif style.type == WD_STYLE_TYPE.TABLE:
                table_styles.append(style.name)

        logger.info(f"Styles listed from: {file_path}")

        return {
            'success': True,
            'styles': all_styles,
            'paragraph_styles': sorted(paragraph_styles),
            'character_styles': sorted(character_styles),
            'table_styles': sorted(table_styles),
            'total_count': len(all_styles)
        }

    except Exception as e:
        logger.error(f"Failed to list styles: {e}", exc_info=True)
        return {'success': False, 'error': f'Failed to list styles: {str(e)}'}
